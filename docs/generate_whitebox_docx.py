from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Tabel 4.14')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = subtitle.add_run('Skenario Pengujian ')
run1.bold = True
run1.font.size = Pt(12)
run1.font.name = 'Times New Roman'
run2 = subtitle.add_run('White Box (Unit Testing)')
run2.bold = True
run2.italic = True
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'

# Define all test data
test_data = [
    # Section A
    {"section": "A. Modul Autentikasi"},
    ("WB-01", "login()", "Identifier atau password kosong", '{ email:"", password:"" }', "422\nValidation Error — field email & password required"),
    ("WB-02", "login()", "Email atau password salah (credential tidak valid)", '{ email:"wrong@mail.com", password:"salah123" }', '422\n"These credentials do not match our records"'),
    ("WB-03", "login()", "Login berhasil sebagai admin/owner/arsitek", '{ email:"admin@mail.com", password:"Password1!" }', "302\nRedirect ke admin.dashboard"),
    ("WB-04", "login()", "Login berhasil sebagai user biasa", '{ email:"user@mail.com", password:"Password1!" }', "302\nRedirect ke user.dashboard"),
    ("WB-05", "logout()", "User melakukan logout", "Session aktif", "302\nSession invalidated, redirect ke /"),
    ("WB-06", "register()", "Semua field kosong", '{ name:"", email:"", phone:"", password:"" }', "422\nValidation Error — semua field required"),
    ("WB-07", "register()", "Email sudah terdaftar (duplikat)", '{ name:"Test", email:"existing@mail.com", phone:"081234567890", password:"Pass1!abc", password_confirmation:"Pass1!abc" }', '422\n"The email has already been taken"'),
    ("WB-08", "register()", "Password tidak memenuhi complexity (tanpa simbol)", '{ name:"Test", email:"new@mail.com", phone:"081234567890", password:"Password1", password_confirmation:"Password1" }', '422\n"The password must contain at least one symbol"'),
    ("WB-09", "register()", "Password confirmation tidak cocok", '{ name:"Test", email:"new@mail.com", phone:"081234567890", password:"Pass1!abc", password_confirmation:"Beda1!abc" }', '422\n"The password confirmation does not match"'),
    ("WB-10", "register()", "Registrasi berhasil, nomor HP diawali 0 → diubah ke 62", '{ name:"Test User", email:"new@mail.com", phone:"081234567890", password:"Pass1!abc", password_confirmation:"Pass1!abc" }', "201\nUser dibuat, phone = 6281234567890, role User di-assign, redirect ke user.dashboard"),
    ("WB-11", "register()", "Role 'User' tidak ditemukan di database", "Data valid, tabel roles kosong", 'User tetap dibuat, error di-log: "Role \'User\' not found"'),
    ("WB-12", "forgotPassword()", "Email kosong", '{ email:"" }', "422\nValidation Error — email required"),
    ("WB-13", "forgotPassword()", "Email valid terdaftar, link terkirim", '{ email:"user@mail.com" }', '302\nRedirect back with status "passwords.sent"'),
    ("WB-14", "resetPassword()", "Token tidak valid / expired", '{ token:"invalid-token", email:"user@mail.com", password:"NewPass1!", password_confirmation:"NewPass1!" }', '302\nRedirect back — "This password reset token is invalid"'),
    ("WB-15", "resetPassword()", "Reset berhasil dengan token valid", '{ token:"valid-token", email:"user@mail.com", password:"NewPass1!", password_confirmation:"NewPass1!" }', "302\nRedirect ke route login dengan status success"),

    # Section B
    {"section": "B. Modul Profil"},
    ("WB-16", "updateProfile()", "Email diubah (isDirty)", '{ name:"Updated", email:"newemail@mail.com" }', "302\nemail_verified_at di-set null, profil tersimpan"),
    ("WB-17", "updateProfile()", "Nama diubah, email tetap", '{ name:"New Name", email:"same@mail.com" }', "302\nProfil diupdate, email_verified_at tetap"),
    ("WB-18", "updatePhoto()", "File bukan gambar (PDF)", "{ profile_photo: file.pdf }", '422\n"The profile photo must be an image"'),
    ("WB-19", "updatePhoto()", "Upload foto valid, foto lama ada → dihapus", "{ profile_photo: valid.jpg } (user sudah punya foto)", "302\nFoto lama dihapus, foto baru tersimpan"),
    ("WB-20", "updatePhoto()", "Upload foto valid, tidak ada foto lama", "{ profile_photo: valid.jpg } (user belum punya foto)", "302\nFoto baru tersimpan"),
    ("WB-21", "deleteAccount()", "Password salah", '{ password:"salahpassword" }', '422\n"The password is incorrect"'),
    ("WB-22", "deleteAccount()", "Password benar, user memiliki foto profil", '{ password:"correctpassword" }', "302\nFoto dihapus, user dihapus, session invalidated, redirect /"),

    # Section C
    {"section": "C. Modul Order (User)"},
    ("WB-23", "createOrder()", "Field wajib kosong", '{ client_type:"", property_type:"", name:"", phone:"", province:"", city:"", district:"", address:"" }', "422\nValidation Error — semua field wajib required"),
    ("WB-24", "createOrder()", "client_type tidak valid (bukan Residensial/Bisnis)", '{ client_type:"Invalid", ... }', '422\n"The selected client type is invalid"'),
    ("WB-25", "createOrder()", "Order berhasil — tipe Residensial, dengan design_type & notes", '{ client_type:"Residensial", property_type:"Rumah", name:"John", phone:"081234567890", province:"Sumut", city:"Medan", district:"Medan Kota", address:"Jl. Test", design_type:["Modern"], notes:"Catatan" }', '200 JSON\n{ success:true, whatsapp_url:"..." }, Order & OrderDetail tersimpan'),
    ("WB-26", "createOrder()", "Order berhasil — tipe Bisnis, tanpa optional fields", '{ client_type:"Bisnis", property_type:"Kantor", name:"Corp", phone:"081234567890", province:"Sumut", city:"Medan", district:"Medan Kota", address:"Jl. Office" }', "200 JSON\n{ success:true }, field opsional null"),
    ("WB-27", "cancelOrder()", "Reason terlalu pendek (<10 karakter)", '{ cancellation_reason:"pendek" }', '422\n"The cancellation reason must be at least 10 characters"'),
    ("WB-28", "cancelOrder()", "Pembatalan berhasil dengan alasan valid", '{ cancellation_reason:"Saya ingin membatalkan pesanan ini" }', "302\nOrderDetail 'cancelled' dibuat, redirect ke index"),
    ("WB-29", "cancelOrder()", "User mencoba membatalkan order milik orang lain", "User ID ≠ Order user_id", "403\nForbidden — Unauthorized"),

    # Section D
    {"section": "D. Modul Pembayaran (Midtrans)"},
    ("WB-30", "createPayment()", "User bukan pemilik order", "Order user_id ≠ Auth id", '403\n"Unauthorized"'),
    ("WB-31", "createPayment()", "Sudah ada pending payment dengan snap_token", "Order milik user, pending payment exists", "200\nView payment dengan snapToken dari payment lama"),
    ("WB-32", "createPayment()", "Belum ada pending, transaksi Midtrans berhasil", "Order milik user, tidak ada pending payment", "200\nView payment dengan snapToken baru"),
    ("WB-33", "createPayment()", "Belum ada pending, transaksi Midtrans gagal", "Midtrans return { success:false }", '302\nRedirect back — "Gagal membuat pembayaran"'),
    ("WB-34", "notification()", "Signature key tidak valid", '{ signature_key:"invalid_hash", order_id:"ORDER-1-xxx", status_code:"200", gross_amount:"500000" }', '403 JSON\n{ message:"Invalid signature" }'),
    ("WB-35", "notification()", "Signature key valid", "Notification data dengan hash SHA512 valid", '200 JSON\n{ message:"OK" }'),
    ("WB-36", "finishPayment()", "transaction_status = settlement", '{ order_id:"ORDER-1-xxx", transaction_status:"settlement" }', "302\nPayment status → success, paid_at di-set"),
    ("WB-37", "finishPayment()", "transaction_status = pending", '{ order_id:"ORDER-1-xxx", transaction_status:"pending" }', "302\nPayment status → pending"),
    ("WB-38", "finishPayment()", "transaction_status = deny/cancel/expire", '{ order_id:"ORDER-1-xxx", transaction_status:"deny" }', "302\nPayment status → failed"),
    ("WB-39", "finishPayment()", "order_id kosong / null", "{ order_id: null }", "302\nRedirect ke user.dashboard"),

    # Section E
    {"section": "E. Modul Feedback"},
    ("WB-40", "storeFeedback()", "User bukan pemilik order", "Order user_id ≠ Auth id", '403\n"Unauthorized"'),
    ("WB-41", "storeFeedback()", "Order belum berstatus completed", "latestDetail.status ≠ completed", '302\n"Anda hanya bisa memberikan feedback untuk pesanan yang sudah selesai"'),
    ("WB-42", "storeFeedback()", "Feedback sudah pernah diberikan", "order→feedback exists", '302\n"Anda sudah memberikan feedback untuk pesanan ini"'),
    ("WB-43", "storeFeedback()", "Rating tidak valid (0 atau 6)", '{ rating:0, review:"Test" }', '422\n"The rating must be at least 1"'),
    ("WB-44", "storeFeedback()", "Feedback berhasil disimpan", '{ rating:5, review:"Sangat memuaskan", would_recommend:true }', '302\n"Terima kasih! Feedback Anda telah disimpan"'),

    # Section F
    {"section": "F. Modul Kategori (Admin)"},
    ("WB-45", "storeCategory()", "Nama kosong", '{ name:"" }', '422\n"The name field is required"'),
    ("WB-46", "storeCategory()", "Nama duplikat", '{ name:"Kategori Existing" }', '422\n"The name has already been taken"'),
    ("WB-47", "storeCategory()", "Data valid, slug auto-generated", '{ name:"Modern Minimalis", description:"Desain modern" }', '302\nKategori tersimpan, slug = "modern-minimalis"'),
    ("WB-48", "updateCategory()", "Nama diubah ke nama yang sudah ada", '{ name:"Nama Kategori Lain" } (duplikat)', '422\n"The name has already been taken"'),
    ("WB-49", "updateCategory()", "Nama tetap sama (unique ignore self)", '{ name:"Nama Sama" } (nama asli)', "302\nKategori diupdate, slug di-regenerate"),
    ("WB-50", "destroyCategory()", "Kategori valid dihapus", "Category ID valid", "302\nKategori dihapus, redirect ke index"),

    # Section G
    {"section": "G. Modul Proyek (Admin)"},
    ("WB-51", "storeProject()", "Tanpa gambar (images kosong)", '{ title:"Proyek A", category_id:1, project_date:"2026-01-01", description:"Desc" }', '422\n"The images field is required"'),
    ("WB-52", "storeProject()", "Data valid dengan gambar, featured_image di-set", '{ title:"Proyek A", category_id:1, project_date:"2026-01-01", description:"Desc", images:[img1,img2], featured_image:0 }', "302\nProyek & gambar tersimpan, index 0 is_featured=true"),
    ("WB-53", "updateProject()", "Update tanpa gambar baru, featured image lama dipilih", "{ ..., existing_featured_image:5 }", "302\nexisting_featured_image di-set is_featured=true"),
    ("WB-54", "updateProject()", "Update dengan gambar baru, index baru di-set featured", "{ ..., images:[newImg], featured_image_index:0 }", "302\nGambar lama is_featured=false, gambar baru is_featured=true"),
    ("WB-55", "updateProject()", "Tidak ada featured dipilih → gambar pertama jadi featured", "{ ..., existing_featured_image:null } (tanpa gambar baru)", "302\nGambar pertama otomatis is_featured=true"),

    # Section H
    {"section": "H. Modul Order (Admin)"},
    ("WB-56", "updateOrder()", "Status tidak valid", '{ status:"unknown_status" }', '422\n"The selected status is invalid"'),
    ("WB-57", "updateOrder()", "Update berhasil tanpa foto", '{ status:"in_progress", progress_details:"Sedang dikerjakan", team_members:[1,2] }', "302\nOrderDetail baru dibuat, notifikasi terkirim"),
    ("WB-58", "updateOrder()", "Update berhasil dengan foto baru", '{ status:"completed", progress_details:"Selesai", new_photos:[photo1.jpg] }', "302\nFoto diupload, OrderDetail dengan photos"),
    ("WB-59", "updateOrder()", "Order tidak punya user (user null)", "order→user = null", "302\nOrderDetail dibuat, notifikasi TIDAK terkirim"),

    # Section I
    {"section": "I. Modul Manajemen User (Admin)"},
    ("WB-60", "indexUsers()", "Pencarian dengan keyword", '{ search:"john" }', 'Hanya user dengan nama/email mengandung "john"'),
    ("WB-61", "indexUsers()", "Sort by kolom yang diizinkan", '{ sort:"name", direction:"asc" }', "200\nUser diurutkan nama ascending"),
    ("WB-62", "indexUsers()", "Sort by kolom TIDAK diizinkan", '{ sort:"password", direction:"asc" }', "200\nSort diabaikan, default created_at desc"),
    ("WB-63", "resetUserPassword()", "Password < 8 karakter", '{ password:"short", password_confirmation:"short" }', '422\n"Password minimal 8 karakter"'),
    ("WB-64", "resetUserPassword()", "Password valid & confirmed", '{ password:"NewPass123", password_confirmation:"NewPass123" }', "302\nPassword di-hash & diupdate, redirect ke index"),

    # Section J
    {"section": "J. Modul Tim (Admin)"},
    ("WB-65", "storeTeamMember()", "Posisi tidak valid (tidak ada di daftar)", '{ name:"John", position:"CEO" }', '422\n"The selected position is invalid"'),
    ("WB-66", "storeTeamMember()", "Data valid dengan foto", '{ name:"John", position:"Head of Design", photo:valid.jpg }', "302\nAnggota tim tersimpan, foto diupload"),
    ("WB-67", "storeTeamMember()", "Data valid tanpa foto", '{ name:"John", position:"Head of Design" }', "302\nAnggota tim tersimpan tanpa photo_path"),
    ("WB-68", "updateTeamMember()", "Upload foto baru, foto lama ada → dihapus", '{ name:"John", position:"Head of Design", photo:new.jpg }', "302\nFoto lama dihapus, foto baru diupload"),
    ("WB-69", "updateTeamMember()", "Remove foto (checkbox)", '{ name:"John", position:"Head of Design", remove_photo:true }', "302\nphoto_path = null, file dihapus"),
    ("WB-70", "destroyTeamMember()", "Anggota tim memiliki foto", "TeamMember dengan photo_path", "302\nFoto dihapus dari storage, anggota dihapus"),
    ("WB-71", "destroyTeamMember()", "Anggota tim tanpa foto", "TeamMember tanpa photo_path", "302\nAnggota dihapus, tidak ada penghapusan file"),
]

# Create table
headers = ["ID", "Unit/Fungsi yang Diuji", "Jalur Logika\n(Branch/Path)", "Data Uji\n(Contoh)", "Output yang\nDiharapkan"]
col_count = 5

table = doc.add_table(rows=1, cols=col_count)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
widths = [Cm(1.8), Cm(3.2), Cm(4.0), Cm(5.0), Cm(4.5)]

# Header row
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = ''
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    # Gray background for header
    shading = hdr_cells[i]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'D9E2F3',
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

# Data rows
for item in test_data:
    if isinstance(item, dict):
        # Section header
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[4])
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(item["section"])
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        # Light blue background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
    else:
        row = table.add_row()
        for i, val in enumerate(item):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            if i == 1:  # Underline function name
                run.underline = True

# Set column widths
for row in table.rows:
    for i, width in enumerate(widths):
        row.cells[i].width = width

# Save
output_path = r'c:\laragon\www\desain-interior-web\docs\white-box-testing.docx'
doc.save(output_path)
print(f"File Word berhasil dibuat: {output_path}")
